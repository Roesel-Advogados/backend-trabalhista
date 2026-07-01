"""Gera a peça em .docx com o timbre do escritório (logo no cabeçalho).

Edite CABECALHO e RODAPE com os dados reais da Roesel Advogados.
A logo deve ficar em app/assets/logo-roesel.jpg
"""
import io
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ASSETS = os.path.join(os.path.dirname(__file__), "..", "assets")
LOGO_PATH = os.path.join(ASSETS, "logo-roesel.jpg")

# ---- Dados do timbre (ajuste ao escritório) ----------------------------
CABECALHO = {
    "nome": "ROESEL ADVOGADOS",
    "linha2": "Advocacia Trabalhista",
}
RODAPE = "Roesel Advogados  ·  OAB/__ nº ____  ·  contato@roesel.adv.br  ·  (34) ____-____"

TEAL = RGBColor(0x12, 0x7D, 0x8E)
CINZA = RGBColor(0x56, 0x61, 0x69)


def _monta_timbre(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.18)   # ~3cm
    section.right_margin = Inches(0.79)  # ~2cm

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        try:
            hp.add_run().add_picture(LOGO_PATH, width=Inches(1.6))
        except Exception:
            pass
    nome_p = header.add_paragraph()
    nome_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = nome_p.add_run(CABECALHO["nome"])
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = TEAL
    sub = header.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(CABECALHO["linha2"])
    rs.font.size = Pt(9)
    rs.font.color.rgb = CINZA

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp.add_run(RODAPE)
    rf.font.size = Pt(8)
    rf.font.color.rgb = CINZA


def gerar_docx(conteudo: str, titulo: str | None = None) -> bytes:
    """Recebe o texto da peça e devolve os bytes do .docx com timbre."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    _monta_timbre(doc)

    if titulo:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(titulo.upper())
        tr.bold = True
        tr.font.size = Pt(13)
        doc.add_paragraph()

    blocos = [b.strip() for b in conteudo.split("\n") if b.strip()]
    for bloco in blocos:
        p = doc.add_paragraph(bloco)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.first_line_indent = Inches(0.5)
        pf.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()